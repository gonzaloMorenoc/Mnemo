"""Tests for src/testplan/ingest.py — HU file ingestion (PDF, DOCX, plain text)."""

from __future__ import annotations

import io

import pytest
import pypdf
import docx as python_docx
from reportlab.pdfgen import canvas as rl_canvas

from src.testplan.ingest import (
    text_from_pdf,
    text_from_docx,
    resolve_hu_from_upload,
)

# ---------------------------------------------------------------------------
# Helpers to build minimal in-memory fixtures
# ---------------------------------------------------------------------------

_SIZE_CAP_BYTES = 5 * 1024 * 1024  # 5 MB — must match ingest.py constant


def _make_pdf_bytes(text: str) -> bytes:
    """Build a minimal single-page PDF with the given text via reportlab."""
    buf = io.BytesIO()
    c = rl_canvas.Canvas(buf)
    c.drawString(72, 720, text)
    c.save()
    return buf.getvalue()


def _make_docx_bytes(text: str) -> bytes:
    """Build a minimal single-paragraph DOCX with the given text."""
    doc = python_docx.Document()
    doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# text_from_pdf
# ---------------------------------------------------------------------------


class TestTextFromPdf:
    def test_extracts_text_from_valid_pdf(self) -> None:
        """text_from_pdf returns the text embedded in a single-page PDF."""
        expected = "Hello PDF World"
        data = _make_pdf_bytes(expected)
        result = text_from_pdf(data)
        assert expected in result

    def test_returns_str(self) -> None:
        data = _make_pdf_bytes("some text")
        assert isinstance(text_from_pdf(data), str)

    def test_empty_pdf_returns_empty_or_whitespace(self) -> None:
        """A blank PDF with no text stream should return empty/whitespace-only."""
        writer = pypdf.PdfWriter()
        writer.add_blank_page(width=612, height=792)
        buf = io.BytesIO()
        writer.write(buf)
        result = text_from_pdf(buf.getvalue())
        assert result.strip() == ""

    def test_invalid_bytes_raise_value_error(self) -> None:
        with pytest.raises(ValueError, match="PDF"):
            text_from_pdf(b"not a pdf at all")


# ---------------------------------------------------------------------------
# text_from_docx
# ---------------------------------------------------------------------------


class TestTextFromDocx:
    def test_extracts_paragraph_text(self) -> None:
        expected = "Historia de usuario: Login"
        data = _make_docx_bytes(expected)
        result = text_from_docx(data)
        assert expected in result

    def test_multiple_paragraphs_joined(self) -> None:
        doc = python_docx.Document()
        doc.add_paragraph("First paragraph.")
        doc.add_paragraph("Second paragraph.")
        buf = io.BytesIO()
        doc.save(buf)

        result = text_from_docx(buf.getvalue())
        assert "First paragraph." in result
        assert "Second paragraph." in result

    def test_returns_str(self) -> None:
        data = _make_docx_bytes("text")
        assert isinstance(text_from_docx(data), str)

    def test_invalid_bytes_raise_value_error(self) -> None:
        with pytest.raises(ValueError, match="DOCX"):
            text_from_docx(b"not a docx file")


# ---------------------------------------------------------------------------
# resolve_hu_from_upload
# ---------------------------------------------------------------------------


class TestResolveHuFromUpload:
    def test_dispatches_pdf(self) -> None:
        data = _make_pdf_bytes("PDF HU text")
        result = resolve_hu_from_upload("historia.pdf", data)
        assert "PDF HU text" in result

    def test_dispatches_docx(self) -> None:
        data = _make_docx_bytes("DOCX HU text")
        result = resolve_hu_from_upload("historia.docx", data)
        assert "DOCX HU text" in result

    def test_dispatches_txt(self) -> None:
        data = "Plain text HU".encode("utf-8")
        result = resolve_hu_from_upload("historia.txt", data)
        assert "Plain text HU" in result

    def test_unsupported_extension_raises(self) -> None:
        with pytest.raises(ValueError, match="extensión"):
            resolve_hu_from_upload("historia.xlsx", b"data")

    def test_oversized_file_raises(self) -> None:
        oversized = b"x" * (_SIZE_CAP_BYTES + 1)
        with pytest.raises(ValueError, match="tamaño"):
            resolve_hu_from_upload("historia.txt", oversized)

    def test_size_at_cap_is_accepted(self) -> None:
        # Exactly at the cap should not raise
        at_cap = b"a" * _SIZE_CAP_BYTES
        # This is a valid .txt file at exactly the size limit
        result = resolve_hu_from_upload("hu.txt", at_cap)
        assert isinstance(result, str)

    def test_case_insensitive_extension(self) -> None:
        data = _make_docx_bytes("UPPER EXT")
        result = resolve_hu_from_upload("HU.DOCX", data)
        assert "UPPER EXT" in result

    def test_txt_returns_decoded_string(self) -> None:
        text = "Cómo usuario quiero iniciar sesión"
        data = text.encode("utf-8")
        result = resolve_hu_from_upload("hu.txt", data)
        assert text in result
